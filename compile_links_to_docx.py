from __future__ import annotations

import argparse
import re
import time
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Iterable
from urllib.parse import urlparse

import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


@dataclass(frozen=True)
class ExtractedPage:
    url: str
    title: str
    subtitle: str | None
    published: str | None
    updated: str | None
    text_blocks: list[str]


def _clean(s: str) -> str:
    s = re.sub(r"\s+", " ", s or "").strip()
    return s


def _is_probably_url(line: str) -> bool:
    try:
        u = urlparse(line.strip())
        return u.scheme in {"http", "https"} and bool(u.netloc)
    except Exception:
        return False


def read_urls(path: str) -> list[str]:
    with open(path, "r", encoding="utf-8") as f:
        lines = [ln.strip() for ln in f.read().splitlines()]
    urls: list[str] = []
    for ln in lines:
        if not ln:
            continue
        if ln.startswith("#"):
            continue
        if _is_probably_url(ln):
            urls.append(ln)
    # de-dupe but keep order
    seen: set[str] = set()
    out: list[str] = []
    for u in urls:
        if u in seen:
            continue
        seen.add(u)
        out.append(u)
    return out


def fetch_html(url: str, *, timeout_s: float = 30.0) -> str:
    headers = {
        "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X) AppleWebKit/537.36 "
        "(KHTML, like Gecko) Chrome/124.0 Safari/537.36"
    }
    r = requests.get(url, headers=headers, timeout=timeout_s)
    r.raise_for_status()
    return r.text


def extract_freespeechproject_entry(html: str, url: str) -> ExtractedPage:
    soup = BeautifulSoup(html, "html.parser")

    # Title
    title = None
    h1 = soup.select_one("h1")
    if h1:
        title = _clean(h1.get_text(" "))
    if not title:
        og = soup.select_one('meta[property="og:title"]')
        if og and og.get("content"):
            title = _clean(og["content"])
    title = title or url

    # The site commonly includes a "First posted ... Last updated ..." line.
    published = None
    updated = None
    subtitle = None

    for p in soup.select("p"):
        t = _clean(p.get_text(" "))
        if not t:
            continue
        if "First posted" in t and "Last updated" in t:
            subtitle = t
            m1 = re.search(r"First posted (.+?) Last updated", t)
            m2 = re.search(r"Last updated (.+)$", t)
            if m1:
                published = _clean(m1.group(1))
            if m2:
                updated = _clean(m2.group(1))
            break

    # Main content: prefer <article>, else fallback to the broad content area.
    root = soup.select_one("article")
    if not root:
        root = soup.select_one("main")
    if not root:
        root = soup.body or soup

    # Remove nav/sidebars/scripts/styles
    for sel in [
        "script",
        "style",
        "noscript",
        "header",
        "footer",
        "nav",
        "aside",
        ".sharedaddy",
        ".jp-relatedposts",
        ".post-tags",
        ".entry-meta",
        ".comment",
        "#comments",
    ]:
        for node in root.select(sel):
            node.decompose()

    # Collect headings + paragraphs + lists in order.
    blocks: list[str] = []
    for el in root.select("h2, h3, h4, p, li"):
        txt = _clean(el.get_text(" "))
        if not txt:
            continue
        # Skip theme tag lists and obvious boilerplate
        if txt.startswith("All Associated Themes"):
            continue
        if txt.startswith("External References"):
            continue
        if txt == title:
            continue
        blocks.append(txt)

    # De-dupe consecutive duplicates
    deduped: list[str] = []
    for b in blocks:
        if deduped and b == deduped[-1]:
            continue
        deduped.append(b)

    return ExtractedPage(
        url=url,
        title=title,
        subtitle=subtitle,
        published=published,
        updated=updated,
        text_blocks=deduped,
    )


def add_page_to_doc(doc: Document, page: ExtractedPage) -> None:
    title_p = doc.add_paragraph()
    title_run = title_p.add_run(page.title)
    title_run.bold = True
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(page.url)

    meta_bits = []
    if page.published:
        meta_bits.append(f"First posted: {page.published}")
    if page.updated:
        meta_bits.append(f"Last updated: {page.updated}")
    if page.subtitle and not meta_bits:
        meta_bits.append(page.subtitle)
    if meta_bits:
        mp = doc.add_paragraph(" | ".join(meta_bits))
        mp.italic = True

    doc.add_paragraph("")  # spacer

    for b in page.text_blocks:
        # Promote "Key Players", "Further Details", "Outcome" etc. to headings if they look like section headers.
        if len(b) <= 80 and (b.isupper() or b.istitle()) and not b.endswith("."):
            doc.add_heading(b, level=2)
        else:
            doc.add_paragraph(b)

    doc.add_page_break()


def compile_urls_to_docx(
    urls: Iterable[str],
    output_docx: str,
    *,
    delay_s: float = 0.4,
    timeout_s: float = 30.0,
) -> None:
    out_path = Path(output_docx)
    if out_path.parent and str(out_path.parent) not in {".", ""}:
        out_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    doc.add_heading("Combined Links Export", level=1)
    doc.add_paragraph(f"Generated: {datetime.now().isoformat(timespec='seconds')}")
    doc.add_paragraph("")

    for i, url in enumerate(urls, start=1):
        try:
            html = fetch_html(url, timeout_s=timeout_s)
            page = extract_freespeechproject_entry(html, url)
            doc.add_heading(f"{i}. {page.title}", level=1)
            add_page_to_doc(doc, page)
        except Exception as e:
            doc.add_heading(f"{i}. (FAILED) {url}", level=1)
            doc.add_paragraph(f"Error: {type(e).__name__}: {e}")
            doc.add_page_break()
        time.sleep(delay_s)

    doc.save(str(out_path))


def main() -> int:
    ap = argparse.ArgumentParser(
        description="Read a txt file of URLs and export combined content to a .docx"
    )
    ap.add_argument("--input", required=True, help="Path to .txt file containing URLs")
    ap.add_argument("--output", required=True, help="Path to output .docx")
    ap.add_argument("--delay", type=float, default=0.4, help="Delay between requests (seconds)")
    ap.add_argument("--timeout", type=float, default=30.0, help="HTTP timeout (seconds)")
    args = ap.parse_args()

    urls = read_urls(args.input)
    if not urls:
        raise SystemExit("No URLs found in input file.")

    compile_urls_to_docx(
        urls,
        args.output,
        delay_s=args.delay,
        timeout_s=args.timeout,
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

